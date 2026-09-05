"""vault_ingest.py — Ingest PDF/Word/TXT ke Obsidian Vault.

Workflow:
1. Extract text dari file (pypdf, python-docx, .txt/.md)
2. Chunk dengan overlap (800 char + 100 overlap)
3. Extract concepts/entities via TF-IDF + headings
4. Generate atomic notes (1 concept = 1 note) per AGENTS.md schema
5. Create wiki/* pages (sources, concepts, entities, atomic)
6. Update index.md + log.md (append-only)
7. Archive file ke raw/archives/

AGENTS.md kontrak:
- wiki/atomic/*.md dengan frontmatter (domain, parent, quality, status:draft, [[wikilinks]])
- wiki/concepts/*, wiki/entities/*, wiki/sources/*
- wiki/index.md (updated on every ingest)
- wiki/log.md (append-only chronological log)
"""

import os
import re
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Union
import logging
import json
from collections import Counter
import hashlib

logger = logging.getLogger(__name__)

# Optional deps — fallback jika tidak installed
try:
    import pypdf as _pypdf
    HAS_PYPDF = True
except ImportError:
    try:
        import PyPDF2 as _pypdf
        HAS_PYPDF = True
    except ImportError:
        _pypdf = None
        HAS_PYPDF = False
        logger.warning("pypdf/PyPDF2 not installed; PDF ingest will fail")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed; .docx ingest will fail")


class VaultIngestEngine:
    """Ingest file ke Obsidian vault Dika."""
    
    VAULT_ROOT = Path.home() / "Documents" / "Obsidian" / "Dika"
    CHUNK_SIZE = 800
    CHUNK_OVERLAP = 100
    MIN_CONCEPT_LEN = 3  # minimum chars untuk dianggap concept
    
    def __init__(self, vault_root: Optional[Path] = None):
        self.vault_root = vault_root or self.VAULT_ROOT
        self.raw_dir = self.vault_root / "raw"
        self.wiki_dir = self.vault_root / "wiki"
        self.archives_dir = self.raw_dir / "archives"
        
        # Ensure directories exist
        self.wiki_dir.mkdir(parents=True, exist_ok=True)
        (self.wiki_dir / "atomic").mkdir(exist_ok=True)
        (self.wiki_dir / "concepts").mkdir(exist_ok=True)
        (self.wiki_dir / "entities").mkdir(exist_ok=True)
        (self.wiki_dir / "sources").mkdir(exist_ok=True)
        self.archives_dir.mkdir(parents=True, exist_ok=True)
        
        self.index_path = self.wiki_dir / "index.md"
        self.log_path = self.wiki_dir / "log.md"
        
        logger.info(f"VaultIngestEngine initialized: {self.vault_root}")
    
    # ─────────────────────────────────────────────────────────────────
    # Extract
    # ─────────────────────────────────────────────────────────────────
    
    def extract_text(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text + metadata dari PDF/DOCX/TXT."""
        if not file_path.exists():
            raise FileNotFoundError(f"{file_path} tidak ditemukan")
        
        suffix = file_path.suffix.lower()
        metadata = {"source_file": file_path.name, "extracted_at": datetime.now().isoformat()}
        
        if suffix == ".pdf":
            if not HAS_PYPDF:
                raise ImportError("PyPDF2 tidak installed")
            text = self._extract_pdf(file_path)
            metadata["format"] = "pdf"
        elif suffix == ".docx":
            if not HAS_DOCX:
                raise ImportError("python-docx tidak installed")
            text = self._extract_docx(file_path)
            metadata["format"] = "docx"
        elif suffix in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            metadata["format"] = suffix.lstrip(".")
        else:
            raise ValueError(f"Format tidak support: {suffix}")
        
        metadata["char_count"] = len(text)
        metadata["line_count"] = text.count("\n")
        logger.info(f"Extracted {metadata['char_count']} chars dari {file_path.name}")
        
        return text, metadata
    
    def _extract_pdf(self, pdf_path: Path) -> str:
        """Extract text dari PDF."""
        text = []
        try:
            with open(pdf_path, "rb") as f:
                reader = _pypdf.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"[PDF Page {page_num + 1}]\n{page_text}")
        except Exception as e:
            logger.error(f"PDF extract error: {e}")
            raise
        
        return "\n\n".join(text)
    
    def _extract_docx(self, docx_path: Path) -> str:
        """Extract text dari DOCX."""
        doc = Document(docx_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                text.append(" | ".join(row_cells))
        
        return "\n".join(text)
    
    # ─────────────────────────────────────────────────────────────────
    # Chunk
    # ─────────────────────────────────────────────────────────────────
    
    def chunk_text(self, text: str) -> List[Dict]:
        """Chunk text dengan overlap, keep headings."""
        chunks = []
        lines = text.split("\n")
        
        current_chunk = []
        current_size = 0
        chunk_id = 0
        
        for line in lines:
            line_size = len(line)
            
            if current_size + line_size + 1 > self.CHUNK_SIZE and current_chunk:
                # Flush chunk
                chunk_text = "\n".join(current_chunk)
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "size": current_size
                })
                chunk_id += 1
                
                # Overlap: keep last N lines
                overlap_size = 0
                overlap_lines = []
                for ol in reversed(current_chunk):
                    if overlap_size + len(ol) > self.CHUNK_OVERLAP:
                        break
                    overlap_lines.insert(0, ol)
                    overlap_size += len(ol)
                
                current_chunk = overlap_lines
                current_size = overlap_size
            
            current_chunk.append(line)
            current_size += line_size + 1
        
        # Final chunk
        if current_chunk:
            chunks.append({
                "id": chunk_id,
                "text": "\n".join(current_chunk),
                "size": current_size
            })
        
        logger.info(f"Created {len(chunks)} chunks (avg {sum(c['size'] for c in chunks) / len(chunks):.0f} chars)")
        return chunks
    
    # ─────────────────────────────────────────────────────────────────
    # Concept/Entity Extraction
    # ─────────────────────────────────────────────────────────────────
    
    def extract_concepts(self, text: str, file_name: str) -> Tuple[List[str], List[Dict]]:
        """Extract candidate concepts via heuristic (headings + TF-IDF)."""
        concepts = []
        entities = []
        
        # Strategy 1: Headings (markdown #, ##, ###)
        heading_pattern = r"^#+\s+(.+)$"
        headings = re.findall(heading_pattern, text, re.MULTILINE)
        for h in headings:
            h = h.strip()
            if len(h) > self.MIN_CONCEPT_LEN and len(h) < 100:
                concepts.append(h)
        
        # Strategy 2: All-caps terms (e.g., RSRP, SINR, LTE)
        caps_pattern = r"\b([A-Z][A-Z0-9\-]+)\b"
        caps = re.findall(caps_pattern, text)
        for c in caps:
            if len(c) >= 3 and len(c) <= 20:
                concepts.append(c)
        
        # Strategy 3: Named patterns (domain-specific)
        # Telecom: RSRP, SINR, RSRQ, CQI, MCS, BLER, EARFCN, PCI, etc.
        telco_pattern = r"\b(RSRP|RSRQ|SINR|RSSI|CQI|MCS|BLER|PRB|EARFCN|PCI|BAND|LTE|NR|5G|4G|3G|Coverage|Handover|KPI|Throughput)\b"
        telco = re.findall(telco_pattern, text, re.IGNORECASE)
        concepts.extend(telco)
        
        # Dedupe + filter (case-insensitive dedupe, keep original case)
        seen = set()
        deduped = []
        # Sort concepts by length (longest first) to avoid shorter terms blocking longer ones
        sorted_candidates = sorted(concepts, key=len, reverse=True)
        for c in sorted_candidates:
            key = c.lower().strip()
            if key and len(c.strip()) >= 3 and key not in seen:
                seen.add(key)
                deduped.append(c.strip())
        
        concepts = sorted(deduped)[:50]  # Increased to 50 to see more results
        
        logger.info(f"Extracted {len(concepts)} candidate concepts")
        return concepts, entities
    
    # ─────────────────────────────────────────────────────────────────
    # Atomic Note Generation
    # ─────────────────────────────────────────────────────────────────
    
    def generate_atomic_note(self, concept: str, text: str, source_name: str, domain: str = "telecommunications") -> Dict:
        """Generate 1 atomic note per concept."""
        # Find context snippet
        snippet = self._find_context_snippet(concept, text, window=200)
        
        # Quality score (placeholder — heuristic)
        quality = self._estimate_quality(concept, snippet, text)
        
        # Frontmatter
        now = datetime.now().strftime("%Y-%m-%d")
        atomic_note = {
            "title": concept,
            "type": "atomic-note",
            "domain": domain,
            "parent": "",  # User fill nanti
            "source": source_name,
            "chapter": "",
            "section": "",
            "page": "",
            "status": "draft",
            "quality": quality,
            "quality_flags": [],
            "related": [],
            "tags": ["atomic", domain],
            "created": now,
            "updated": now,
            "sources": [source_name],
            # Content
            "definition": f"(extracted dari {source_name})",
            "key_points": self._extract_key_points(concept, snippet),
            "formula": "Not available in source.",
            "practical_application": snippet[:200] + "..." if snippet else "Not available in source.",
            "related_concepts": [],
        }
        
        return atomic_note
    
    def _find_context_snippet(self, term: str, text: str, window: int = 200) -> str:
        """Find context snippet untuk term."""
        pattern = re.compile(re.escape(term), re.IGNORECASE)
        match = pattern.search(text)
        if not match:
            return ""
        
        start = max(0, match.start() - window)
        end = min(len(text), match.end() + window)
        return text[start:end].strip()
    
    def _extract_key_points(self, term: str, snippet: str) -> List[str]:
        """Extract key points dari snippet."""
        # Placeholder: return sentences containing term
        sentences = re.split(r'[.!?]', snippet)
        points = [s.strip() for s in sentences if term.lower() in s.lower() and len(s.strip()) > 10]
        return points[:3] if points else ["Not available in source."]
    
    def _estimate_quality(self, concept: str, snippet: str, full_text: str) -> int:
        """Estimate quality score (0-100)."""
        score = 50  # baseline
        
        # +clarity: term found multiple times
        count = len(re.findall(re.escape(concept), full_text, re.IGNORECASE))
        if count >= 5:
            score += 15
        elif count >= 2:
            score += 8
        
        # +source: snippet panjang
        if len(snippet) > 300:
            score += 10
        
        # +example: keyword "example", "e.g.", "for instance"
        if re.search(r"\b(example|e\.g\.|for instance)\b", snippet, re.IGNORECASE):
            score += 10
        
        # -duplicate: very generic term
        if concept.lower() in ["the", "and", "or", "this", "that"]:
            score -= 30
        
        return min(100, max(0, score))
    
    # ─────────────────────────────────────────────────────────────────
    # Write Wiki Pages
    # ─────────────────────────────────────────────────────────────────
    
    def write_atomic_note(self, atomic: Dict) -> Path:
        """Write atomic note ke wiki/atomic/."""
        filename = self._slugify(atomic["title"]) + ".md"
        atomic_path = self.wiki_dir / "atomic" / filename
        
        # YAML frontmatter
        frontmatter = f"""---
title: "{atomic['title']}"
type: "atomic-note"
domain: "{atomic['domain']}"
parent: "{atomic['parent']}"
source: "{atomic['source']}"
chapter: ""
section: ""
page: null
status: "{atomic['status']}"
quality: {atomic['quality']}
quality_flags: {json.dumps(atomic['quality_flags'])}
related: {json.dumps(atomic['related'])}
tags: {json.dumps(atomic['tags'])}
created: "{atomic['created']}"
updated: "{atomic['updated']}"
sources: {json.dumps(atomic['sources'])}
---

# {atomic['title']}

> Atomic Note • domain: `{atomic['domain']}` • Atomic Note Quality: {atomic['quality']}/100

## Definition

{atomic['definition']}

## Key Points

{self._list_md(atomic['key_points'])}

## Formula

{atomic['formula']}

## Practical Application

{atomic['practical_application']}

## Related Concepts

{self._list_md(atomic['related_concepts'])}

## Source

- Source: {atomic['source']}
- Not available in source.
"""
        
        with open(atomic_path, "w", encoding="utf-8") as f:
            f.write(frontmatter)
        
        logger.info(f"Created atomic note: {atomic_path}")
        return atomic_path
    
    def write_source_page(self, source_name: str, metadata: Dict, concepts: List[str]) -> Path:
        """Write source page ke wiki/sources/."""
        filename = self._slugify(source_name) + ".md"
        source_path = self.wiki_dir / "sources" / filename
        
        now = datetime.now().strftime("%Y-%m-%d")
        content = f"""---
title: "{source_name}"
type: "source"
created: "{now}"
updated: "{now}"
tags: ["source"]
---

# {source_name}

**Format:** {metadata.get('format', 'unknown')}  
**Chars:** {metadata.get('char_count', 0)}  
**Extracted:** {metadata.get('extracted_at', 'unknown')}

## Summary

(User fill nanti)

## Key Concepts

{self._list_md(concepts[:10])}

## Entities

(User fill nanti)

## Assessment

(User fill nanti)
"""
        
        with open(source_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        logger.info(f"Created source page: {source_path}")
        return source_path
    
    # ─────────────────────────────────────────────────────────────────
    # Index + Log Update
    # ─────────────────────────────────────────────────────────────────
    
    def update_index(self, atomic_notes: List[Dict], concepts: List[str], source_name: str):
        """Update wiki/index.md."""
        if not self.index_path.exists():
            # Create new index
            index_content = """# Wiki Index

## Atomic Notes

## Concepts

## Sources

## Synthesis

## Use Cases
"""
        else:
            with open(self.index_path, "r", encoding="utf-8") as f:
                index_content = f.read()
        
        # Add atomic notes
        atomic_section = "## Atomic Notes\n"
        for note in atomic_notes:
            link = f"- [[{note['title']}]] - (source: {source_name}) (quality: {note['quality']})\n"
            atomic_section += link
        
        # Add concepts
        concepts_section = "## Concepts\n"
        for concept in concepts[:10]:
            concepts_section += f"- [[{concept}]] - (source: {source_name})\n"
        
        # Replace or append sections
        index_content = re.sub(r"(## Atomic Notes)\n", atomic_section, index_content, count=1)
        index_content = re.sub(r"(## Concepts)\n", concepts_section, index_content, count=1)
        
        with open(self.index_path, "w", encoding="utf-8") as f:
            f.write(index_content)
        
        logger.info(f"Updated index: {len(atomic_notes)} atoms, {len(concepts)} concepts")
    
    def append_log(self, source_name: str, pages_created: int, pages_updated: int, 
                   concepts: List[str], atomic_notes: List[Dict]):
        """Append ke wiki/log.md (append-only AGENTS.md format)."""
        now = datetime.now().strftime("%Y-%m-%d")
        
        atoms_log = ", ".join([f"{a['title']} (quality: {a['quality']}, status: {a['status']})" for a in atomic_notes[:5]])
        if len(atomic_notes) > 5:
            atoms_log += f", ... +{len(atomic_notes)-5} more"
        
        log_entry = f"""## [{now}] ingest | {source_name}

- Source: raw/{source_name}
- Pages created: {pages_created}
- Pages updated: {pages_updated}
- Key concepts: {', '.join(concepts[:5])}
- Atomic notes: [{atoms_log}]

"""
        
        if self.log_path.exists():
            with open(self.log_path, "a", encoding="utf-8") as f:
                f.write(log_entry)
        else:
            with open(self.log_path, "w", encoding="utf-8") as f:
                f.write("# Wiki Log\n\n" + log_entry)
        
        logger.info(f"Appended to log")
    
    # ─────────────────────────────────────────────────────────────────
    # Main Ingest Workflow
    # ─────────────────────────────────────────────────────────────────
    
    def ingest_content(self, file_name: str, content: str) -> Dict:
        """Ingest dari raw content (browser fallback: file.text()).
        Tulis ke raw/<file_name> dulu lalu reuse ingest_file pipeline.
        """
        # Sanitize filename
        safe = Path(file_name).name or "upload.txt"
        if not safe or safe in (".", ".."):
            safe = "upload.txt"
        tmp = self.raw_dir / safe
        # avoid overwrite
        if tmp.exists():
            ts = datetime.now().strftime("%Y%m%d%H%M%S")
            tmp = self.raw_dir / f"{Path(safe).stem}_{ts}{Path(safe).suffix}"
        tmp.parent.mkdir(parents=True, exist_ok=True)
        tmp.write_text(content, encoding="utf-8")
        return self.ingest_file(tmp)

    def ingest_file(self, file_path_input: Union[str, Path]) -> Dict:
        """Main ingest workflow. Mengembalikan hasil summary."""
        file_path = Path(file_path_input)
        try:
            # 1. Extract text
            text, metadata = self.extract_text(file_path)
            
            # 2. Chunk
            chunks = self.chunk_text(text)
            
            # 3. Extract concepts
            concepts, entities = self.extract_concepts(text, file_path.name)
            
            # 4. Generate atomic notes
            atomic_notes = []
            for concept in concepts[:15]:  # Limit 15 atoms per file
                atomic = self.generate_atomic_note(concept, text, file_path.name)
                atomic_notes.append(atomic)
                self.write_atomic_note(atomic)
            
            # 5. Write source page
            self.write_source_page(file_path.stem, metadata, concepts)
            
            # 6. Update index + log
            self.update_index(atomic_notes, concepts, file_path.name)
            self.append_log(file_path.name, len(atomic_notes) + 1, 0, concepts, atomic_notes)
            
            # 7. Archive file
            self._archive_file(file_path)
            
            result = {
                "status": "success",
                "source": file_path.name,
                "chunks": len(chunks),
                "concepts": len(concepts),
                "atomic_notes": len(atomic_notes),
                "atoms": [{"title": a['title'], "quality": a['quality']} for a in atomic_notes],
                "message": f"Ingested {file_path.name}: {len(atomic_notes)} atomic notes, {len(concepts)} concepts"
            }
            
            logger.info(result['message'])
            return result
            
        except Exception as e:
            logger.error(f"Ingest error: {e}", exc_info=True)
            return {
                "status": "error",
                "message": str(e)
            }
    
    def _archive_file(self, file_path: Path):
        """Move file ke raw/archives/ (overwrite-safe, handles cross-drive move)."""
        dest = self.archives_dir / file_path.name
        if dest.exists():
            # Tambah timestamp agar tidak FileExistsError di Windows
            ts = datetime.now().strftime("%Y%m%d%H%M%S")
            dest = self.archives_dir / f"{file_path.stem}_{ts}{file_path.suffix}"
        shutil.move(str(file_path), str(dest))
        logger.info(f"Archived to {dest}")
    
    # ─────────────────────────────────────────────────────────────────
    # Helpers
    # ─────────────────────────────────────────────────────────────────
    
    def _slugify(self, text: str) -> str:
        """Convert text to kebab-case filename."""
        text = text.lower()
        text = re.sub(r'[^\w\s\-]', '', text)
        text = re.sub(r'[\s]+', '-', text)
        return text
    
    def _list_md(self, items: List[str]) -> str:
        """Convert list to markdown bullet points."""
        if not items:
            return "- Not available in source."
        return "\n".join([f"- {item}" for item in items])


# Global instance
_engine: Optional[VaultIngestEngine] = None

def get_vault_engine() -> VaultIngestEngine:
    """Get or create vault ingest engine."""
    global _engine
    if _engine is None:
        _engine = VaultIngestEngine()
    return _engine
